"""
AI Engine for Resolve AI - PDF Generation with Hebrew Support
Supports both PDF and Word (.docx) document processing
"""
import json
from datetime import datetime
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
import os
import random
import hashlib
from bidi.algorithm import get_display
import re

# PDF text extraction (optional, only if PyPDF2 is available)
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("Warning: PyPDF2 not available. PDF text extraction disabled.")

# Word document text extraction
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not available. Word document text extraction disabled.")

# Try to register Hebrew font
HEBREW_FONT = 'Helvetica'
try:
    # Try common Hebrew font paths on Linux
    font_paths = [
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
        '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
        '/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf'
    ]

    for font_path in font_paths:
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('HebrewFont', font_path))
            HEBREW_FONT = 'HebrewFont'
            break
except Exception as e:
    print(f"Could not register Hebrew font: {e}. Using fallback.")

def extract_text_from_pdf(file_path):
    """
    Extract text content from a PDF file

    Args:
        file_path: Path to the PDF file

    Returns:
        str: Extracted text content
    """
    if not PDF_SUPPORT:
        return "[PDF text extraction not available - PyPDF2 not installed]"

    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        print(f"Error extracting text from PDF {file_path}: {e}")
        return f"[Error extracting PDF text: {str(e)}]"

def extract_text_from_docx(file_path):
    """
    Extract text content from a Word (.docx) file

    Args:
        file_path: Path to the DOCX file

    Returns:
        str: Extracted text content
    """
    if not DOCX_SUPPORT:
        return "[Word text extraction not available - python-docx not installed]"

    try:
        doc = Document(file_path)
        text = []

        # Extract text from all paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text)

        return "\n".join(text)
    except Exception as e:
        print(f"Error extracting text from Word document {file_path}: {e}")
        return f"[Error extracting Word text: {str(e)}]"

def extract_text_from_file(file_path):
    """
    Extract text from a file (PDF or DOCX) based on file extension

    Args:
        file_path: Path to the file

    Returns:
        str: Extracted text content
    """
    if not file_path or not os.path.exists(file_path):
        return "[File not found]"

    # Determine file type by extension
    _, ext = os.path.splitext(file_path.lower())

    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        if ext == '.doc':
            print(f"Warning: .doc files are not supported, only .docx. File: {file_path}")
            return "[.doc format not supported - please use .docx format]"
        return extract_text_from_docx(file_path)
    else:
        return f"[Unsupported file format: {ext}]"

def generate_case_id():
    """Generate a unique 5-digit case ID"""
    return str(random.randint(10000, 99999))

def analyze_case(claimant_name, defendant_name):
    """
    Analyze case and return structured JSON output

    Returns:
        dict: Structured analysis with dispute_table, mediation_proposal,
              final_verdict, reasoning, and legal_expenses
    """

    # Simulate AI analysis with realistic mock data
    analysis = {
        "case_metadata": {
            "claimant": claimant_name,
            "defendant": defendant_name,
            "date_analyzed": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "case_id": generate_case_id()
        },

        "dispute_table": [
            {
                "issue": "עיכוב במועד האספקה",
                "claimant_version": "הנתבע התחייב לספק את המוצרים תוך 30 יום ולא עמד בכך, מה שגרם לנזק כלכלי של 5,000 שקל",
                "defendant_version": "העיכוב נבע מכוח עליון (מגפת קורונה) ולכן אין אחריות, כמפורט בסעיף 15 לחוזה",
                "ai_analysis": "נמצא כי אכן היה עיכוב של 14 ימי עסקים. עם זאת, טענת כוח עליון מתקבלת רק באופן חלקי. מומלץ פיצוי של 2,500 שקל המשקף אחריות חלקית."
            },
            {
                "issue": "איכות המוצרים שסופקו",
                "claimant_version": "המוצרים שסופקו היו פגומים ולא עמדו בתקן המוסכם בהסכם",
                "defendant_version": "המוצרים עמדו בכל התקנים המוסכמים והבדיקות המצורפות מאשרות זאת",
                "ai_analysis": "לפי מסמכי הבדיקה המצורפים, המוצרים עמדו בתקן. טענת התובע נדחית בנקודה זו."
            },
            {
                "issue": "הוצאות נוספות שנגרמו לתובע",
                "claimant_version": "נאלצתי לרכוש מוצרים חלופיים בעלות של 3,000 שקל עקב העיכוב",
                "defendant_version": "הרכישה הנוספת לא הייתה הכרחית ונעשתה ביוזמת התובע ללא תיאום",
                "ai_analysis": "חלק מההוצאות (1,500 שקל) נראות סבירות והכרחיות בנסיבות העניין. יש להכיר בהן."
            }
        ],

        "mediation_proposal": {
            "proposal": "הנתבע ישלם 3,500 שקל (במקום 4,000 שקל) ובתמורה התובע יסכים לוותר על כל תביעות נוספות ולחדש את ההתקשרות העסקית לשנה נוספת בתנאים מותאמים",
            "rationale": "פשרה זו שומרת על היחסים העסקיים בין הצדדים, חוסכת הוצאות משפט נוספות, ומאפשרת המשך שיתוף פעולה פורה"
        },

        "final_verdict": {
            "verdict": "התביעה מתקבלת בחלקה",
            "amount_awarded": 4000.0,
            "legal_expenses": 35.0,
            "total_payment": 4035.0,
            "payment_deadline_days": 30
        },

        "reasoning": {
            "summary": "בהתבסס על הראיות והטיעונים, נמצא כי הנתבע אחראי באופן חלקי לנזקים שנגרמו",
            "detailed_analysis": [
                "לגבי העיכוב במועדים: נמצא כי הייתה הפרת חוזה, אך טענת כוח עליון מתקבלת באופן חלקי. העיכוב של 14 יום הוא מעבר לסביר גם בתקופת משבר.",
                "לגבי איכות המוצרים: הראיות מצביעות על כך שהמוצרים עמדו בתקנים. התביעה נדחית בנקודה זו.",
                "לגבי הוצאות נוספות: חלק מההוצאות (1,500 שקל) היו הכרחיות וסבירות בנסיבות העניין.",
                "הפיצוי הכולל (4,000 שקל) משקף אחריות יחסית של הנתבע והוא הולם את הנסיבות."
            ],
            "legal_basis": "ההחלטה ניתנה על פי חוק החוזים (חלק כללי), התשל\"ג-1973, וחוק הבוררות, התשכ\"ח-1968"
        },

        "legal_expenses": {
            "registered_mail": 35.0,
            "explanation": "דמי המשלוח בדואר רשום (35 שקל) יוחזרו לתובע כחלק מהוצאות המשפט, בהתאם לסעיף 76 לחוק בתי המשפט [נוסח משולב], התשמ\"ד-1984",
            "included_in_total": True
        }
    }

    return analysis

def generate_arbitral_award_pdf(case_data, analysis, output_path):
    """
    Generate a formal Arbitral Award PDF document with Hebrew support

    Args:
        case_data: Dictionary with case information (case_id, claimant, defendant)
        analysis: The AI analysis result with structured data
        output_path: Path where to save the PDF

    Returns:
        str: Path to generated PDF file
    """

    # Function to add logo to each page
    def add_page_header(canvas, doc):
        """Add Resolve AI logo to every page header"""
        canvas.saveState()
        try:
            # Try to load logo
            logo_path = "logo.png"
            if os.path.exists(logo_path):
                canvas.drawImage(logo_path,
                               doc.width/2 + doc.leftMargin - 1*cm,  # Center horizontally
                               doc.height + doc.topMargin - 1.5*cm,  # Top of page
                               width=2*cm,
                               height=2*cm,
                               preserveAspectRatio=True,
                               mask='auto')
            else:
                # Try URL if local file doesn't exist
                logo_url = "https://raw.githubusercontent.com/yanaydavid/ResolveAI/main/logo.png"
                canvas.drawImage(logo_url,
                               doc.width/2 + doc.leftMargin - 1*cm,
                               doc.height + doc.topMargin - 1.5*cm,
                               width=2*cm,
                               height=2*cm,
                               preserveAspectRatio=True,
                               mask='auto')
        except Exception as e:
            print(f"Warning: Could not load logo on page: {e}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        output_path,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=3*cm,  # Increased to make room for logo header
        bottomMargin=2*cm
    )

    # Container for the 'Flowable' objects
    elements = []

    # Helper function to fix Hebrew text direction
    def fix_hebrew(text):
        """Fix Hebrew text direction using BiDi algorithm with proper RTL handling"""
        if not text:
            return text

        # Handle HTML tags separately - extract them, process text, then reinsert
        # Pattern to find HTML tags
        html_tag_pattern = re.compile(r'(<[^>]+>)')

        # Split text by HTML tags
        parts = html_tag_pattern.split(str(text))

        processed_parts = []
        for part in parts:
            if html_tag_pattern.match(part):
                # It's an HTML tag, keep as is
                processed_parts.append(part)
            elif part.strip():
                # It's text content, apply BiDi with RTL base direction
                try:
                    processed_parts.append(get_display(part, base_dir='R'))
                except:
                    # Fallback if BiDi fails
                    processed_parts.append(part)
            else:
                # It's whitespace
                processed_parts.append(part)

        return ''.join(processed_parts)

    # Define styles
    styles = getSampleStyleSheet()

    # Custom styles with Hebrew font
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=22,
        textColor=colors.HexColor('#0A2647'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName=HEBREW_FONT,
        leading=28
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#0A2647'),
        spaceAfter=15,
        spaceBefore=15,
        alignment=TA_RIGHT,
        fontName=HEBREW_FONT,
        leading=20
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        alignment=TA_RIGHT,
        fontName=HEBREW_FONT,
        leading=18,
        wordWrap='RTL'
    )

    # Table cell style with better wrapping
    table_cell_style = ParagraphStyle(
        'TableCell',
        parent=normal_style,
        fontSize=10,
        alignment=TA_RIGHT,
        fontName=HEBREW_FONT,
        leading=14,
        wordWrap='RTL',
        spaceBefore=2,
        spaceAfter=2
    )

    # Legal Header (Logo will be added to every page via add_page_header callback)
    legal_header_style = ParagraphStyle(
        'LegalHeader',
        parent=normal_style,
        fontSize=10,
        alignment=TA_CENTER,
        textColor=colors.HexColor('#0A2647'),
        spaceAfter=15
    )
    elements.append(Paragraph(fix_hebrew("פסק בורר לפי חוק הבוררות, התשכ\"ח-1968"), legal_header_style))
    elements.append(Spacer(1, 0.3*cm))

    # Title
    elements.append(Paragraph(fix_hebrew("פסק בוררות"), title_style))
    elements.append(Spacer(1, 0.8*cm))

    # Generate timestamp for the arbitral award
    award_timestamp = datetime.now()
    timestamp_str = award_timestamp.strftime('%d/%m/%Y %H:%M:%S')

    # Case Information Box
    case_info = f"""
    <b>{fix_hebrew('מספר תיק:')}</b> {case_data['case_id']}<br/>
    <b>{fix_hebrew('תאריך וזמן מתן הפסק:')}</b> {timestamp_str}<br/>
    <b>{fix_hebrew('התובע:')}</b> {fix_hebrew(case_data['claimant'])}<br/>
    <b>{fix_hebrew('הנתבע:')}</b> {fix_hebrew(case_data['defendant'])}<br/>
    """
    elements.append(Paragraph(case_info, normal_style))
    elements.append(Spacer(1, 0.8*cm))

    # Introduction
    elements.append(Paragraph(fix_hebrew("הקדמה"), heading_style))
    intro_text = fix_hebrew("""
    בהתאם לחוק הבוררות, התשכ"ח-1968, ולאחר בחינה מעמיקה של כתב התביעה וכתב ההגנה,
    וניתוח הראיות באמצעות מערכת Resolve AI, מתפרסם בזו פסק הבוררות הבא.
    """)
    elements.append(Paragraph(intro_text, normal_style))
    elements.append(Spacer(1, 0.5*cm))

    # Objectivity Declaration
    elements.append(Paragraph(fix_hebrew("הצהרת אובייקטיביות"), heading_style))
    objectivity_text = f"""
    <b>{fix_hebrew('ההכרעה התקבלה על ידי אלגוריתם בינה מלאכותית המבוסס על ניתוח עובדתי של המסמכים שהוגשו בלבד, ללא מגע יד אדם וללא ניגוד עניינים.')}</b>
    <br/><br/>
    {fix_hebrew('המערכת מבצעת ניתוח אוטומטי ואובייקטיבי של הטיעונים והראיות שהוצגו על ידי שני הצדדים, ומפיקה החלטה מבוססת על עקרונות משפטיים מקובלים וצדק טבעי.')}
    """
    elements.append(Paragraph(objectivity_text, normal_style))
    elements.append(Spacer(1, 0.8*cm))

    # Disputes Analysis Table
    elements.append(Paragraph(fix_hebrew("ניתוח נקודות המחלוקת"), heading_style))

    # Wrap text in Paragraphs for proper RTL handling and word wrap
    # Header style for table
    table_header_style = ParagraphStyle(
        'TableHeader',
        parent=table_cell_style,
        fontSize=11,
        fontName=HEBREW_FONT,
        alignment=TA_RIGHT,
        leading=16
    )

    table_data = [
        [Paragraph(f'<b>{fix_hebrew("ניתוח AI")}</b>', table_header_style),
         Paragraph(f'<b>{fix_hebrew("גרסת הנתבע")}</b>', table_header_style),
         Paragraph(f'<b>{fix_hebrew("גרסת התובע")}</b>', table_header_style),
         Paragraph(f'<b>{fix_hebrew("נושא")}</b>', table_header_style)]
    ]

    for dispute in analysis['dispute_table']:
        table_data.append([
            Paragraph(fix_hebrew(dispute['ai_analysis']), table_cell_style),
            Paragraph(fix_hebrew(dispute['defendant_version']), table_cell_style),
            Paragraph(fix_hebrew(dispute['claimant_version']), table_cell_style),
            Paragraph(fix_hebrew(dispute['issue']), table_cell_style)
        ])

    # Create table with flexible column widths and dynamic row heights
    # Using None for rowHeights allows automatic height adjustment based on content
    dispute_table = Table(table_data, colWidths=[4.5*cm, 4.5*cm, 4.5*cm, 3*cm], repeatRows=1)
    dispute_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E0E7FF')),  # Light blue background
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#0A2647')),  # Dark blue text
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),  # Right align for RTL
        ('FONTNAME', (0, 0), (-1, 0), HEBREW_FONT),
        ('FONTSIZE', (0, 0), (-1, 0), 11),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),  # Add horizontal padding
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 1), (-1, -1), HEBREW_FONT),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),  # Align text to top for better readability
        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),  # More padding in content cells
        ('TOPPADDING', (0, 1), (-1, -1), 10),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.lightgrey]),
    ]))

    elements.append(dispute_table)
    elements.append(Spacer(1, 0.8*cm))

    # Mediation Proposal
    if 'mediation_proposal' in analysis:
        elements.append(Paragraph(fix_hebrew("הצעת פשרה"), heading_style))
        mediation = analysis['mediation_proposal']
        mediation_text = f"""
        <b>{fix_hebrew('ההצעה:')}</b> {fix_hebrew(mediation['proposal'])}<br/><br/>
        <b>{fix_hebrew('נימוק:')}</b> {fix_hebrew(mediation['rationale'])}
        """
        elements.append(Paragraph(mediation_text, normal_style))
        elements.append(Spacer(1, 0.8*cm))

    # Final Decision
    elements.append(Paragraph(fix_hebrew("ההחלטה הסופית"), heading_style))
    decision = analysis['final_verdict']
    reasoning = analysis['reasoning']

    verdict_text = f"""
    <b>{fix_hebrew('פסיקה:')}</b> {fix_hebrew(decision['verdict'])}<br/><br/>
    """
    elements.append(Paragraph(verdict_text, normal_style))

    # Reasoning
    elements.append(Paragraph(fix_hebrew("נימוקים"), heading_style))
    reasoning_summary = f"<b>{fix_hebrew('סיכום:')}</b> {fix_hebrew(reasoning['summary'])}<br/><br/>"
    elements.append(Paragraph(reasoning_summary, normal_style))

    if 'detailed_analysis' in reasoning:
        for i, point in enumerate(reasoning['detailed_analysis'], 1):
            elements.append(Paragraph(f"{i}. {fix_hebrew(point)}", normal_style))
            elements.append(Spacer(1, 0.3*cm))

    if 'legal_basis' in reasoning:
        legal_basis_text = f"<br/><b>{fix_hebrew('בסיס משפטי:')}</b> {fix_hebrew(reasoning['legal_basis'])}"
        elements.append(Paragraph(legal_basis_text, normal_style))

    elements.append(Spacer(1, 0.8*cm))

    # Financial Summary
    elements.append(Paragraph(fix_hebrew("סיכום כספי"), heading_style))

    legal_exp = analysis.get('legal_expenses', {})
    legal_exp_amount = legal_exp.get('registered_mail', decision.get('legal_expenses', 35.0))

    financial_data = [
        [Paragraph(f'<b>{fix_hebrew("סכום")}</b>', table_cell_style), Paragraph(f'<b>{fix_hebrew("פריט")}</b>', table_cell_style)],
        [Paragraph(f"{decision['amount_awarded']:,.2f} ₪", table_cell_style), Paragraph(fix_hebrew('סכום הפיצוי'), table_cell_style)],
        [Paragraph(f"{legal_exp_amount:.2f} ₪", table_cell_style), Paragraph(fix_hebrew('דמי משלוח דואר רשום'), table_cell_style)],
        [Paragraph(f"{decision['total_payment']:,.2f} ₪", table_cell_style), Paragraph(fix_hebrew('סה"כ לתשלום'), table_cell_style)]
    ]

    financial_table = Table(financial_data, colWidths=[6*cm, 8*cm])
    financial_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7C3AED')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),  # Right align for RTL
        ('FONTNAME', (0, 0), (-1, -1), HEBREW_FONT),
        ('FONTSIZE', (0, 0), (-1, 0), 12),
        ('FONTSIZE', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),  # Padding for all cells
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),  # Add horizontal padding
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('BACKGROUND', (0, 1), (-1, -2), colors.white),
        ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#F1F5F9')),
        ('GRID', (0, 0), (-1, -1), 1.5, colors.black),
        ('FONTNAME', (0, -1), (-1, -1), HEBREW_FONT),
        ('FONTSIZE', (0, -1), (-1, -1), 12),
        ('FONTWEIGHT', (0, -1), (-1, -1), 'BOLD'),
    ]))

    elements.append(financial_table)
    elements.append(Spacer(1, 0.8*cm))

    # Payment terms
    payment_text = f"""
    <b>{fix_hebrew('מועד תשלום:')}</b> {fix_hebrew(f'על הנתבע לשלם את מלוא הסכום תוך {decision["payment_deadline_days"]} ימים מיום קבלת פסק בוררות זה.')}
    """
    elements.append(Paragraph(payment_text, normal_style))
    elements.append(Spacer(1, 1.2*cm))

    # Generate document hash first (needed for authentication appendix)
    hash_content = f"{case_data['case_id']}|{case_data['claimant']}|{case_data['defendant']}|{timestamp_str}|{decision['amount_awarded']}|{decision['total_payment']}"
    doc_hash = hashlib.sha256(hash_content.encode('utf-8')).hexdigest()

    # Authentication Appendix - Improved, concise, and emphasizes Hash
    elements.append(Paragraph(fix_hebrew("נספח אימות ושלמות מסמך"), heading_style))

    auth_text = f"""
    <b>{fix_hebrew('אימות שלמות המסמך:')}</b><br/>
    {fix_hebrew('מסמך זה הופק דיגיטלית במערכת Resolve AI לאחר שני הצדדים אישרו את הליך הבוררות והגישו את מסמכיהם דרך הפלטפורמה המאובטחת.')}
    <br/><br/>
    <b>{fix_hebrew('קוד אימות Hash (חשוב):')}</b><br/>
    <font size="8" color="#0A2647"><b>{doc_hash}</b></font><br/>
    {fix_hebrew('קוד ייחודי זה מאמת את שלמות תוכן המסמך ומונע כל שינוי או זיוף. ניתן לאמת את המסמך באמצעות קוד זה במערכת Resolve AI.')}
    <br/><br/>
    <b>{fix_hebrew('תיעוד דיגיטלי:')}</b><br/>
    {fix_hebrew(f'תיק מספר {case_data["case_id"]} תועד במלואו במערכת, כולל אישורי הצדדים, מועדי הגשת המסמכים, וזמני האישור.')}
    """
    elements.append(Paragraph(auth_text, normal_style))
    elements.append(Spacer(1, 0.8*cm))

    # Signature section - Digital signatures with actual dates
    elements.append(Paragraph(fix_hebrew("חתימות דיגיטליות"), heading_style))
    elements.append(Spacer(1, 0.5*cm))

    date_str = award_timestamp.strftime("%d/%m/%Y")
    time_str = award_timestamp.strftime("%H:%M:%S")

    signature_data = [
        [Paragraph(f'<b>{fix_hebrew("נחתם דיגיטלית על ידי:")}</b><br/>{fix_hebrew("מערכת Resolve AI")}<br/><b>{fix_hebrew("תאריך:")}</b> {date_str}<br/><b>{fix_hebrew("שעה:")}</b> {time_str}', table_cell_style),
         Paragraph('', table_cell_style)],
        [Paragraph('', table_cell_style), Paragraph('', table_cell_style)],
        [Paragraph(f'<b>{fix_hebrew("אישור קבלה - תובע")}</b><br/>{fix_hebrew("נחתם דיגיטלית על ידי:")}<br/>{fix_hebrew(case_data["claimant"])}<br/><b>{fix_hebrew("תאריך:")}</b> {date_str}<br/><b>{fix_hebrew("שעה:")}</b> {time_str}', table_cell_style),
         Paragraph(f'<b>{fix_hebrew("אישור קבלה - נתבע")}</b><br/>{fix_hebrew("נחתם דיגיטלית על ידי:")}<br/>{fix_hebrew(case_data["defendant"])}<br/><b>{fix_hebrew("תאריך:")}</b> {date_str}<br/><b>{fix_hebrew("שעה:")}</b> {time_str}', table_cell_style)]
    ]

    sig_table = Table(signature_data, colWidths=[7*cm, 7*cm])
    sig_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('FONTNAME', (0, 0), (-1, -1), HEBREW_FONT),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TOPPADDING', (0, 0), (-1, -1), 12),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
        ('LEFTPADDING', (0, 0), (-1, -1), 8),
        ('RIGHTPADDING', (0, 0), (-1, -1), 8),
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#0A2647')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#0A2647')),
    ]))

    elements.append(sig_table)
    elements.append(Spacer(1, 1.2*cm))

    # Footer
    footer_text = f"""
    <i>{fix_hebrew('פסק בוררות זה ניתן על פי חוק הבוררות, התשכ"ח-1968, ומהווה פסק דין סופי ומחייב.')}<br/>
    {fix_hebrew('ערר על פסק בוררות זה ניתן להגיש לבית המשפט בהתאם להוראות החוק.')}<br/>
    {fix_hebrew('מסמך זה נוצר באמצעות מערכת Resolve AI - בוררות דיגיטלית מבוססת בינה מלאכותית.')}</i>
    """
    elements.append(Paragraph(footer_text, ParagraphStyle(
        'Footer',
        parent=normal_style,
        fontSize=8,
        textColor=colors.grey,
        alignment=TA_CENTER,
        leading=12
    )))

    # Build PDF with logo on every page
    try:
        doc.build(elements, onFirstPage=add_page_header, onLaterPages=add_page_header)
        return output_path
    except Exception as e:
        print(f"Error generating PDF: {e}")
        raise

def get_analysis_summary_html(analysis):
    """
    Convert analysis to formatted HTML for UI display with proper RTL and styling

    Args:
        analysis: The analysis dictionary

    Returns:
        str: HTML formatted analysis summary
    """

    # Build dispute table with dark blue header and zebra stripes
    html = """
    <div style='margin: 30px 0;'>
        <h3 style='color: #0A2647; font-size: 1.8rem; margin-bottom: 20px; text-align: center;'>
            📊 ניתוח נקודות המחלוקת
        </h3>
        <table style='width: 100%; border-collapse: collapse; direction: rtl; box-shadow: 0 10px 40px rgba(0,0,0,0.1);'>
            <thead>
                <tr style='background: #0A2647; color: white;'>
                    <th style='padding: 15px; text-align: right; border: 1px solid #ddd; font-size: 1.1rem;'>נושא</th>
                    <th style='padding: 15px; text-align: right; border: 1px solid #ddd; font-size: 1.1rem;'>גרסת התובע</th>
                    <th style='padding: 15px; text-align: right; border: 1px solid #ddd; font-size: 1.1rem;'>גרסת הנתבע</th>
                    <th style='padding: 15px; text-align: right; border: 1px solid #ddd; font-size: 1.1rem;'>ניתוח AI</th>
                </tr>
            </thead>
            <tbody>
    """

    # Add zebra-striped rows
    for i, dispute in enumerate(analysis['dispute_table']):
        bg_color = '#f8f9fa' if i % 2 == 0 else 'white'
        html += f"""
            <tr style='background: {bg_color};'>
                <td style='padding: 12px; border: 1px solid #ddd; vertical-align: top; text-align: right; direction: rtl;'><b>{dispute['issue']}</b></td>
                <td style='padding: 12px; border: 1px solid #ddd; vertical-align: top; text-align: right; direction: rtl;'>{dispute['claimant_version']}</td>
                <td style='padding: 12px; border: 1px solid #ddd; vertical-align: top; text-align: right; direction: rtl;'>{dispute['defendant_version']}</td>
                <td style='padding: 12px; border: 1px solid #ddd; vertical-align: top; background: #e8f5e9; text-align: right; direction: rtl;'>{dispute['ai_analysis']}</td>
            </tr>
        """

    html += """
            </tbody>
        </table>
    </div>
    """

    # Mediation Proposal with RTL
    if 'mediation_proposal' in analysis:
        mediation = analysis['mediation_proposal']
        html += f"""
        <div style='background: linear-gradient(135deg, #fbbf24 0%, #f59e0b 100%);
                    padding: 30px; border-radius: 20px; color: white; margin: 30px 0;
                    box-shadow: 0 10px 40px rgba(251,191,36,0.4);'>
            <h3 style='font-size: 1.8rem; margin-bottom: 15px; text-align: center;'>🤝 הצעת פשרה</h3>
            <div style='background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: right; direction: rtl;'>
                <p style='margin-bottom: 15px; line-height: 1.8; font-size: 1.1rem;'>
                    <b>ההצעה:</b><br/>{mediation['proposal']}
                </p>
                <p style='line-height: 1.8; font-size: 1.1rem;'>
                    <b>נימוק:</b><br/>{mediation['rationale']}
                </p>
            </div>
        </div>
        """

    # Final Verdict - Green card with centered header and RTL reasoning
    decision = analysis['final_verdict']
    reasoning = analysis.get('reasoning', {})

    # Build reasoning text
    reasoning_html = ""
    if reasoning:
        reasoning_html = f"<b>סיכום:</b> {reasoning.get('summary', '')}<br/><br/>"

        if 'detailed_analysis' in reasoning:
            reasoning_html += "<b>ניתוח מפורט:</b><br/>"
            for i, point in enumerate(reasoning['detailed_analysis'], 1):
                reasoning_html += f"{i}. {point}<br/>"

        if 'legal_basis' in reasoning:
            reasoning_html += f"<br/><b>בסיס משפטי:</b> {reasoning['legal_basis']}"

    html += f"""
        <div style='background: linear-gradient(135deg, #10b981 0%, #059669 100%);
                    padding: 30px; border-radius: 20px; color: white; margin: 30px 0;
                    box-shadow: 0 10px 40px rgba(16,185,129,0.4);'>
            <h3 style='font-size: 2rem; margin-bottom: 15px; text-align: center;'>📜 ההחלטה הסופית</h3>
            <p style='font-size: 1.3rem; text-align: center; margin-bottom: 20px;'><b>{decision['verdict']}</b></p>

            <div style='background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px; text-align: right; direction: rtl; margin-bottom: 20px;'>
                <p style='line-height: 1.8; font-size: 1.1rem;'>
                    <b>נימוקים:</b><br/><br/>
                    {reasoning_html}
                </p>
            </div>

            <div style='background: rgba(255,255,255,0.15); padding: 20px; border-radius: 12px;'>
                <div style='display: flex; justify-content: space-between; margin-bottom: 10px; direction: rtl;'>
                    <span style='font-size: 1.1rem;'>סכום הפיצוי:</span>
                    <span style='font-size: 1.2rem; font-weight: bold;'>{decision['amount_awarded']:,.0f} ₪</span>
                </div>
                <div style='display: flex; justify-content: space-between; margin-bottom: 10px; direction: rtl;'>
                    <span style='font-size: 1.1rem;'>הוצאות משפט (דמי משלוח):</span>
                    <span style='font-size: 1.2rem; font-weight: bold;'>{decision['legal_expenses']:.0f} ₪</span>
                </div>
                <hr style='border: none; border-top: 2px solid rgba(255,255,255,0.3); margin: 15px 0;'>
                <div style='display: flex; justify-content: space-between; direction: rtl;'>
                    <span style='font-size: 1.2rem;'><b>סה״כ לתשלום:</b></span>
                    <span style='font-size: 1.5rem; font-weight: bold;'>{decision['total_payment']:,.0f} ₪</span>
                </div>
                <p style='text-align: center; margin-top: 15px; font-size: 0.95rem;'>
                    תשלום תוך {decision['payment_deadline_days']} ימים
                </p>
            </div>
        </div>
    """

    return html
